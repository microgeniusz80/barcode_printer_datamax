import os
from docx import Document
from docx.shared import Cm, Pt
from tempfile import NamedTemporaryFile

# === Config ===
barcode_data = "55667788"
barcode_font = "Code 128"
text_font = "Calibri"

# === Create document ===
doc = Document()

# Set sticker size to 5.5cm x 2cm
section = doc.sections[0]
section.page_width = Cm(5.5)
section.page_height = Cm(2)
section.top_margin = Cm(0.1)
section.bottom_margin = Cm(0.1)
section.left_margin = Cm(0.2)
section.right_margin = Cm(0.2)

# === One paragraph ===
p = doc.add_paragraph()
p.alignment = 1  # Centered

# Barcode line
run1 = p.add_run(barcode_data)
run1.font.name = barcode_font
run1.font.size = Pt(24)

# Line break + number
run1.add_break()
run2 = p.add_run(barcode_data)
run2.font.name = text_font
run2.font.size = Pt(10)

# === Save and print ===
with NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
    doc_path = tmp_file.name
    doc.save(doc_path)

os.startfile(doc_path, "print")
