import os
from barcode import Code128
from barcode.writer import ImageWriter
from PIL import Image
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from tempfile import NamedTemporaryFile
from datetime import datetime

# === Config ===
barcode_data = "508020005"
ward = "Ward 8"
line2_text = "NORSYUHADA BINTI MOHD SABRI"
ic_number = "010722120354"
CROP_KEEP = 0.3  # portion of barcode height to keep
BARCODE_WIDTH_CM = 4.0  # barcode print width (adjust as needed)
BARCODE_HEIGHT_CM = 1 # barcode print height

# === Generate barcode and crop ===
barcode_img_path = "barcode.png"
barcode_cropped_path = "barcode_cropped.png"
barcode = Code128(barcode_data, writer=ImageWriter())
barcode.save(barcode_img_path.replace(".png", ""))

img = Image.open(barcode_img_path)
width, height = img.size
cropped_img = img.crop((0, 0, width, int(height * CROP_KEEP)))
cropped_img.save(barcode_cropped_path)

# === Create Word document ===
doc = Document()
section = doc.sections[0]
section.page_width = Cm(5.5)
section.page_height = Cm(2)
section.top_margin = Cm(0)
section.bottom_margin = Cm(0)
section.left_margin = Cm(0)
section.right_margin = Cm(0)

# === Paragraph with tight spacing ===
p = doc.add_paragraph()
p.alignment = 1  # Center
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 0.8  # tighter lines

# === Barcode image ===
run = p.add_run()
run.add_picture(barcode_cropped_path, width=Cm(BARCODE_WIDTH_CM), height=Cm(BARCODE_HEIGHT_CM))

# === First line: barcode data + ward ===
run.add_break()
line1 = f"{barcode_data} ({ward})"
text1 = p.add_run(line1)
text1.font.name = "Calibri"
text1.font.size = Pt(8)
text1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

# === Second line: patient name ===
text1.add_break()
text2 = p.add_run(line2_text)
text2.font.name = "Calibri"
text2.font.size = Pt(8)
text2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

# === Third line: date/time + IC ===
now = datetime.now()
date_str = now.strftime("%d/%m")
time_str = now.strftime("%H:%M")  # no seconds
line3 = f"{date_str} - {time_str} - IC: {ic_number}"
text2.add_break()
text3 = p.add_run(line3)
text3.font.name = "Calibri"
text3.font.size = Pt(8)
text3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')


# === Save and print ===
with NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
    doc_path = tmp_file.name
    doc.save(doc_path)

os.startfile(doc_path, "print")
