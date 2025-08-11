import os
import sqlite3
from datetime import datetime
from pathlib import Path
import tkinter as tk
from tkinter import ttk, messagebox

# Sticker libs
from barcode import Code128
from barcode.writer import ImageWriter
from PIL import Image
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from tempfile import NamedTemporaryFile

DB_PATH = "lis.db"

# =========================
# Database Setup / Helpers
# =========================
def init_db():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS patients (
            barcode_data TEXT PRIMARY KEY,
            ward         TEXT,
            name         TEXT,
            ic_number    TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS tests (
            id           INTEGER PRIMARY KEY AUTOINCREMENT,
            barcode_data TEXT,
            test_name    TEXT,
            test_date    TEXT,   -- ISO 'YYYY-MM-DD'
            result       TEXT,
            FOREIGN KEY (barcode_data) REFERENCES patients(barcode_data)
        )
    """)
    conn.commit()
    conn.close()

def upsert_patient(barcode_data, ward, name, ic_number):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO patients (barcode_data, ward, name, ic_number)
        VALUES (?, ?, ?, ?)
        ON CONFLICT(barcode_data) DO UPDATE SET ward=excluded.ward, name=excluded.name, ic_number=excluded.ic_number
    """, (barcode_data, ward, name, ic_number))
    conn.commit()
    conn.close()

def get_patient(barcode_data):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT barcode_data, ward, name, ic_number FROM patients WHERE barcode_data=?", (barcode_data,))
    row = cur.fetchone()
    conn.close()
    return row

def add_test(barcode_data, test_name, test_date, result):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("INSERT INTO tests (barcode_data, test_name, test_date, result) VALUES (?,?,?,?)",
                (barcode_data, test_name, test_date, result))
    conn.commit()
    conn.close()

def list_tests(barcode_data):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("""
        SELECT id, test_name, test_date, result
        FROM tests
        WHERE barcode_data=?
        ORDER BY test_date DESC, id DESC
    """, (barcode_data,))
    rows = cur.fetchall()
    conn.close()
    return rows

# =========================
# Sticker Printing (your working code, wrapped)
# =========================
def print_sticker(barcode_data, ward, name, ic_number,
                  CROP_KEEP=0.3, BARCODE_WIDTH_CM=4.0, BARCODE_HEIGHT_CM=1.0):
    # === Generate barcode and crop ===
    barcode_img_path = "barcode.png"
    barcode_cropped_path = "barcode_cropped.png"
    barcode = Code128(barcode_data, writer=ImageWriter())
    barcode.save(barcode_img_path.replace(".png", ""))  # creates barcode.png

    img = Image.open(barcode_img_path)
    width, height = img.size
    cropped_img = img.crop((0, 0, width, int(height * CROP_KEEP)))
    cropped_img.save(barcode_cropped_path)

    # === Create Word document (5.5cm x 2cm, no margins) ===
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(5.5)
    section.page_height = Cm(2)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)

    # === Paragraph with tight spacing ===
    p = doc.add_paragraph()
    p.alignment = 1  # center
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 0.8

    # === Barcode image ===
    run = p.add_run()
    run.add_picture(barcode_cropped_path, width=Cm(BARCODE_WIDTH_CM), height=Cm(BARCODE_HEIGHT_CM))

    # === First line: bold barcode number + normal ward ===
    run.add_break()
    bold_run = p.add_run(barcode_data)
    bold_run.font.name = "Calibri"
    bold_run.font.size = Pt(10)
    bold_run.bold = True
    bold_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    normal_run = p.add_run(f" ({ward})")
    normal_run.font.name = "Calibri"
    normal_run.font.size = Pt(8)
    normal_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # === Second line: patient name ===
    normal_run.add_break()
    text2 = p.add_run(name)
    text2.font.name = "Calibri"
    text2.font.size = Pt(8)
    text2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # === Third line: date - time - IC ===
    now = datetime.now()
    date_str = now.strftime("%d/%m")
    time_str = now.strftime("%H:%M")
    line3 = f"{date_str} - {time_str} - IC: {ic_number}"
    text2.add_break()
    text3 = p.add_run(line3)
    text3.font.name = "Calibri"
    text3.font.size = Pt(8)
    text3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # === Save and print ===
    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        doc_path = tmp_file.name
        doc.save(doc_path)

    os.startfile(doc_path, "print")

# =========================
# GUI (Tkinter)
# =========================
class LISApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Mini LIS - Patients & Tests")
        self.geometry("820x520")
        self.resizable(False, False)

        # Patient frame
        frm_pat = ttk.LabelFrame(self, text="Patient")
        frm_pat.place(x=10, y=10, width=390, height=220)

        ttk.Label(frm_pat, text="Barcode:").grid(row=0, column=0, sticky="e", padx=6, pady=6)
        ttk.Label(frm_pat, text="Ward:").grid(row=1, column=0, sticky="e", padx=6, pady=6)
        ttk.Label(frm_pat, text="Name:").grid(row=2, column=0, sticky="e", padx=6, pady=6)
        ttk.Label(frm_pat, text="IC:").grid(row=3, column=0, sticky="e", padx=6, pady=6)

        self.e_barcode = ttk.Entry(frm_pat, width=30)
        self.e_ward    = ttk.Entry(frm_pat, width=30)
        self.e_name    = ttk.Entry(frm_pat, width=30)
        self.e_ic      = ttk.Entry(frm_pat, width=30)
        self.e_barcode.grid(row=0, column=1, padx=6, pady=6)
        self.e_ward.grid(row=1, column=1, padx=6, pady=6)
        self.e_name.grid(row=2, column=1, padx=6, pady=6)
        self.e_ic.grid(row=3, column=1, padx=6, pady=6)

        btn_save = ttk.Button(frm_pat, text="Save/Update Patient", command=self.save_patient)
        btn_load = ttk.Button(frm_pat, text="Load by Barcode", command=self.load_patient)
        btn_print = ttk.Button(frm_pat, text="Print Sticker", command=self.print_sticker_btn)
        btn_save.grid(row=4, column=0, padx=6, pady=6, sticky="ew")
        btn_load.grid(row=4, column=1, padx=6, pady=6, sticky="ew")
        btn_print.grid(row=5, column=0, columnspan=2, padx=6, pady=6, sticky="ew")

        # Test frame
        frm_test = ttk.LabelFrame(self, text="Add Test")
        frm_test.place(x=10, y=240, width=390, height=260)

        ttk.Label(frm_test, text="Test Name:").grid(row=0, column=0, sticky="e", padx=6, pady=6)
        ttk.Label(frm_test, text="Date (YYYY-MM-DD):").grid(row=1, column=0, sticky="e", padx=6, pady=6)
        ttk.Label(frm_test, text="Result:").grid(row=2, column=0, sticky="ne", padx=6, pady=6)

        self.e_test_name = ttk.Entry(frm_test, width=30)
        self.e_test_date = ttk.Entry(frm_test, width=30)
        self.e_result = tk.Text(frm_test, width=28, height=5)
        self.e_test_name.grid(row=0, column=1, padx=6, pady=6)
        self.e_test_date.grid(row=1, column=1, padx=6, pady=6)
        self.e_result.grid(row=2, column=1, padx=6, pady=6)

        if not self.e_test_date.get():
            self.e_test_date.insert(0, datetime.now().strftime("%Y-%m-%d"))

        btn_add_test = ttk.Button(frm_test, text="Add Test", command=self.add_test_btn)
        btn_add_test.grid(row=3, column=0, columnspan=2, padx=6, pady=6, sticky="ew")

        # Tests list frame
        frm_list = ttk.LabelFrame(self, text="Patient Tests")
        frm_list.place(x=410, y=10, width=400, height=490)
        self.tests_tree = ttk.Treeview(frm_list, columns=("id","name","date","result"), show="headings", height=20)
        self.tests_tree.heading("id", text="ID")
        self.tests_tree.heading("name", text="Test")
        self.tests_tree.heading("date", text="Date")
        self.tests_tree.heading("result", text="Result")
        self.tests_tree.column("id", width=40, anchor="center")
        self.tests_tree.column("name", width=120)
        self.tests_tree.column("date", width=80, anchor="center")
        self.tests_tree.column("result", width=140)
        self.tests_tree.pack(fill="both", expand=True, padx=6, pady=6)

        btn_refresh = ttk.Button(frm_list, text="Refresh", command=self.refresh_tests)
        btn_refresh.pack(padx=6, pady=(0,6), fill="x")

    # ----- GUI actions -----
    def save_patient(self):
        b = self.e_barcode.get().strip()
        w = self.e_ward.get().strip()
        n = self.e_name.get().strip()
        ic = self.e_ic.get().strip()
        if not b:
            messagebox.showerror("Error", "Barcode is required.")
            return
        upsert_patient(b, w, n, ic)
        messagebox.showinfo("Saved", "Patient saved/updated.")

    def load_patient(self):
        b = self.e_barcode.get().strip()
        if not b:
            messagebox.showerror("Error", "Enter a barcode to load.")
            return
        rec = get_patient(b)
        if not rec:
            messagebox.showwarning("Not found", "No patient with that barcode.")
            return
        _, w, n, ic = rec
        self.e_ward.delete(0, tk.END); self.e_ward.insert(0, w or "")
        self.e_name.delete(0, tk.END); self.e_name.insert(0, n or "")
        self.e_ic.delete(0, tk.END); self.e_ic.insert(0, ic or "")
        self.refresh_tests()

    def add_test_btn(self):
        b = self.e_barcode.get().strip()
        if not b:
            messagebox.showerror("Error", "Save/load a patient first.")
            return
        tname = self.e_test_name.get().strip()
        tdate = self.e_test_date.get().strip()
        result = self.e_result.get("1.0", "end").strip()
        if not tname or not tdate:
            messagebox.showerror("Error", "Test name and date are required.")
            return
        add_test(b, tname, tdate, result)
        messagebox.showinfo("Added", "Test added.")
        self.refresh_tests()

    def refresh_tests(self):
        for row in self.tests_tree.get_children():
            self.tests_tree.delete(row)
        b = self.e_barcode.get().strip()
        if not b:
            return
        rows = list_tests(b)
        for rid, name, date, result in rows:
            self.tests_tree.insert("", "end", values=(rid, name, date, result))

    def print_sticker_btn(self):
        b = self.e_barcode.get().strip()
        w = self.e_ward.get().strip()
        n = self.e_name.get().strip()
        ic = self.e_ic.get().strip()
        if not b or not n:
            messagebox.showerror("Error", "Barcode and Name are required.")
            return
        # Call your proven printing routine
        print_sticker(b, w, n, ic)
        messagebox.showinfo("Print", "Sticker sent to printer.")

# ========= main ========
if __name__ == "__main__":
    init_db()
    app = LISApp()
    app.mainloop()
