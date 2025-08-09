import os
from barcode import Code128
from barcode.writer import ImageWriter
from docx import Document
from docx.shared import Cm, Inches, Pt
from tempfile import NamedTemporaryFile

# === Configuration ===
barcode_data = "1900642"
barcode_image_path = "barcode.png"  # Temporary image file
label_width_cm = 5.5
label_height_cm = 2.0

# === Step 1: Generate barcode image ===
barcode = Code128(barcode_data, writer=ImageWriter())
barcode.save("barcode")  # Saves as 'barcode.png' by default

# === Step 2: Create Word document ===
doc = Document()

# Set page size and margins
section = doc.sections[0]
section.page_width = Cm(label_width_cm)
section.page_height = Cm(label_height_cm)
section.top_margin = Cm(0.1)
section.bottom_margin = Cm(0.1)
section.left_margin = Cm(0.2)
section.right_margin = Cm(0.2)

# Add barcode image
doc.add_picture(barcode_image_path, width=Cm(5))  # Adjust width to fit nicely

# Add the number below barcode (centered)
p = doc.add_paragraph()
run = p.add_run(barcode_data)
run.font.size = Pt(10)
p.alignment = 1  # Center

# === Step 3: Save and print ===
with NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
    doc_path = tmp_file.name
    doc.save(doc_path)

print(f"✅ Document created at: {doc_path}")

# === Step 4: Print (Windows only) ===
try:
    os.startfile(doc_path, "print")
except Exception as e:
    print(f"⚠️ Printing failed: {e}")
