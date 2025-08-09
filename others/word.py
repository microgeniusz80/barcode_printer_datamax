import os
from docx import Document
from docx.shared import Cm
from tempfile import NamedTemporaryFile

# Create a temporary .docx file
doc = Document()

# Set page size to 5.5cm x 2cm
sections = doc.sections
for section in sections:
    section.page_width = Cm(5.5)
    section.page_height = Cm(2)
    section.top_margin = Cm(0.1)
    section.bottom_margin = Cm(0.1)
    section.left_margin = Cm(0.2)
    section.right_margin = Cm(0.2)

# Add content (centered)
p = doc.add_paragraph()
p.alignment = 1  # Center
run = p.add_run("HELLO WORLD")
run.bold = True

# Save to a temporary path
with NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
    doc_path = tmp_file.name
    doc.save(doc_path)

# Send to printer (opens with Word or default .docx handler and prints)
print(f"🖨️ Sending to printer: {doc_path}")
os.startfile(doc_path, "print")
