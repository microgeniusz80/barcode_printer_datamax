import os
from barcode import Code128
from barcode.writer import ImageWriter
from PIL import Image
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from tempfile import NamedTemporaryFile

# === Config ===
barcode_data = "941224085918"
line2_text = "Rabiatul Adawiyah Binti Abdul Samad"
barcode_img = "barcode.png"
barcode_cropped = "barcode_cropped.png"

# === Generate barcode and crop bottom 40% to remove text ===
barcode = Code128(barcode_data, writer=ImageWriter())
barcode.save(barcode_img.replace(".png", ""))

img = Image.open(barcode_img)
width, height = img.size
cropped_img = img.crop((0, 0, width, int(height * 0.6)))
cropped_img.save(barcode_cropped)

# === Create Word document ===
doc = Document()
section = doc.sections[0]
section.page_width = Cm(5.5)
section.page_height = Cm(2)
section.top_margin = Cm(0)
section.bottom_margin = Cm(0)
section.left_margin = Cm(0)
section.right_margin = Cm(0)

# === One tightly packed paragraph ===
p = doc.add_paragraph()
p.alignment = 1  # Center
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.0

# === Barcode image ===
run = p.add_run()
run.add_picture(barcode_cropped, width=Cm(5.0), height=Cm(1.0))

# === First line of text ===
run.add_break()
text1 = p.add_run(barcode_data)
text1.font.name = "Calibri"
text1.font.size = Pt(8)
text1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

# === Second line of text ===
text1.add_break()
text2 = p.add_run(line2_text)
text2.font.name = "Calibri"
text2.font.size = Pt(8)
text2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

# === Save and print ===
with NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
    doc_path = tmp_file.name
    doc.save(doc_path)

os.startfile(doc_path, "print")
