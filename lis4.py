import os
import sqlite3
from datetime import datetime
import tkinter as tk
from tkinter import ttk, messagebox

# Sticker libs (used by your print function)
from barcode import Code128
from barcode.writer import ImageWriter
from PIL import Image
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from tempfile import NamedTemporaryFile

DB_PATH = "lis.db"

TEST_CHOICES = [
    "GSH",
    "GXM",
    "DCT",
    "RBC phenotyping",
    "blood grouping",
]

# =========================
# Database Setup / Helpers
# =========================
def init_db():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("""
        CREATE TABLE IF NOT EXISTS patients (
            barcode_data TEXT PRIMARY KEY,
            ward         TEXT,
            name         TEXT,
            ic_number    TEXT
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS tests (
            id           INTEGER PRIMARY KEY AUTOINCREMENT,
            barcode_data TEXT,
            test_name    TEXT,
            test_date    TEXT,   -- ISO 'YYYY-MM-DD'
            result       TEXT,
            price        REAL DEFAULT 0,
            FOREIGN KEY (barcode_data) REFERENCES patients(barcode_data)
        )
    """)
    cur.execute("""
        CREATE TABLE IF NOT EXISTS test_prices (
            test_name TEXT PRIMARY KEY,
            price     REAL NOT NULL
        )
    """)
    # Ensure all test choices exist in the price table (default 0.0)
    for t in TEST_CHOICES:
        cur.execute("INSERT OR IGNORE INTO test_prices (test_name, price) VALUES (?, 0.0)", (t,))
    conn.commit()
    conn.close()

def get_current_price(test_name):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT price FROM test_prices WHERE test_name=?", (test_name,))
    row = cur.fetchone()
    conn.close()
    return float(row[0]) if row else 0.0

def get_all_prices():
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT test_name, price FROM test_prices")
    data = dict(cur.fetchall())
    conn.close()
    return data

def set_all_prices(price_dict):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    for name, price in price_dict.items():
        cur.execute("UPDATE test_prices SET price=? WHERE test_name=?", (price, name))
    conn.commit()
    conn.close()

def upsert_patient(barcode_data, ward, name, ic_number):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO patients (barcode_data, ward, name, ic_number)
        VALUES (?, ?, ?, ?)
        ON CONFLICT(barcode_data) DO UPDATE SET
            ward=excluded.ward, name=excluded.name, ic_number=excluded.ic_number
    """, (barcode_data, ward, name, ic_number))
    conn.commit()
    conn.close()

def get_patient(barcode_data):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("SELECT barcode_data, ward, name, ic_number FROM patients WHERE barcode_data=?", (barcode_data,))
    row = cur.fetchone()
    conn.close()
    return row

def add_test(barcode_data, test_name, test_date, result, price_snapshot):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO tests (barcode_data, test_name, test_date, result, price)
        VALUES (?,?,?,?,?)
    """, (barcode_data, test_name, test_date, result, price_snapshot))
    conn.commit()
    conn.close()

def update_test(test_id, test_name, test_date, result, price_snapshot=None):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    if price_snapshot is None:
        cur.execute("""
            UPDATE tests SET test_name=?, test_date=?, result=? WHERE id=?
        """, (test_name, test_date, result, test_id))
    else:
            # When test_name changes, also update the price snapshot
        cur.execute("""
            UPDATE tests SET test_name=?, test_date=?, result=?, price=? WHERE id=?
        """, (test_name, test_date, result, price_snapshot, test_id))
    conn.commit()
    conn.close()

def delete_test(test_id):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("DELETE FROM tests WHERE id=?", (test_id,))
    conn.commit()
    conn.close()

def list_tests(barcode_data):
    conn = sqlite3.connect(DB_PATH)
    cur = conn.cursor()
    cur.execute("""
        SELECT id, test_name, test_date, result, price
        FROM tests
        WHERE barcode_data=?
        ORDER BY test_date DESC, id DESC
    """, (barcode_data,))
    rows = cur.fetchall()
    conn.close()
    return rows

# =========================
# Sticker Printing (your proven DOCX flow)
# =========================
def print_sticker(barcode_data, ward, name, ic_number,
                  CROP_KEEP=0.3, BARCODE_WIDTH_CM=4.0, BARCODE_HEIGHT_CM=1.0):
    # === Generate barcode and crop ===
    barcode_img_path = "barcode.png"
    barcode_cropped_path = "barcode_cropped.png"
    barcode = Code128(barcode_data, writer=ImageWriter())
    barcode.save(barcode_img_path.replace(".png", ""))  # creates barcode.png

    img = Image.open(barcode_img_path)
    width, height = img.size
    cropped_img = img.crop((0, 0, width, int(height * CROP_KEEP)))
    cropped_img.save(barcode_cropped_path)

    # === Create Word document (5.5cm x 2cm, no margins) ===
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(5.5)
    section.page_height = Cm(2)
    section.top_margin = Cm(0)
    section.bottom_margin = Cm(0)
    section.left_margin = Cm(0)
    section.right_margin = Cm(0)

    # === Paragraph with tight spacing ===
    p = doc.add_paragraph()
    p.alignment = 1  # center
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 0.8

    # === Barcode image ===
    run = p.add_run()
    run.add_picture(barcode_cropped_path, width=Cm(BARCODE_WIDTH_CM), height=Cm(BARCODE_HEIGHT_CM))

    # === First line: bold barcode number + normal ward ===
    run.add_break()
    bold_run = p.add_run(barcode_data)
    bold_run.font.name = "Calibri"
    bold_run.font.size = Pt(10)
    bold_run.bold = True
    bold_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    normal_run = p.add_run(f" ({ward})")
    normal_run.font.name = "Calibri"
    normal_run.font.size = Pt(8)
    normal_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # === Second line: patient name ===
    normal_run.add_break()
    text2 = p.add_run(name)
    text2.font.name = "Calibri"
    text2.font.size = Pt(8)
    text2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # === Third line: date - time - IC ===
    now = datetime.now()
    date_str = now.strftime("%d/%m")
    time_str = now.strftime("%H:%M")
    line3 = f"{date_str} - {time_str} - IC: {ic_number}"
    text2.add_break()
    text3 = p.add_run(line3)
    text3.font.name = "Calibri"
    text3.font.size = Pt(8)
    text3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # === Save and print ===
    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        doc_path = tmp_file.name
        doc.save(doc_path)

    os.startfile(doc_path, "print")

# =========================
# GUI (Tkinter) – previous layout
# =========================
class LISApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Mini LIS - Patients & Tests")
        self.geometry("1020x600")  # a bit wider to accommodate the wider list
        self.resizable(False, False)
        self.selected_test_id = None

        # Patient frame (top-left)
        frm_pat = ttk.LabelFrame(self, text="Patient")
        frm_pat.place(x=10, y=10, width=460, height=230)

        ttk.Label(frm_pat, text="Barcode:").grid(row=0, column=0, sticky="e", padx=6, pady=6)
        ttk.Label(frm_pat, text="Ward:").grid(row=1, column=0, sticky="e", padx=6, pady=6)
        ttk.Label(frm_pat, text="Name:").grid(row=2, column=0, sticky="e", padx=6, pady=6)
        ttk.Label(frm_pat, text="IC:").grid(row=3, column=0, sticky="e", padx=6, pady=6)

        self.e_barcode = ttk.Entry(frm_pat, width=36)
        self.e_ward    = ttk.Entry(frm_pat, width=36)
        self.e_name    = ttk.Entry(frm_pat, width=36)
        self.e_ic      = ttk.Entry(frm_pat, width=36)
        self.e_barcode.grid(row=0, column=1, padx=6, pady=6)
        self.e_ward.grid(row=1, column=1, padx=6, pady=6)
        self.e_name.grid(row=2, column=1, padx=6, pady=6)
        self.e_ic.grid(row=3, column=1, padx=6, pady=6)

        btn_save   = ttk.Button(frm_pat, text="Save/Update Patient", command=self.save_patient)
        btn_load   = ttk.Button(frm_pat, text="Load by Barcode", command=self.load_patient)
        btn_print  = ttk.Button(frm_pat, text="Print Sticker", command=self.print_sticker_btn)
        btn_prices = ttk.Button(frm_pat, text="Set Test Prices", command=self.open_prices_dialog)
        btn_save.grid(row=4, column=0, padx=6, pady=4, sticky="ew")
        btn_load.grid(row=4, column=1, padx=6, pady=4, sticky="ew")
        btn_print.grid(row=5, column=0, padx=6, pady=4, sticky="ew")
        btn_prices.grid(row=5, column=1, padx=6, pady=4, sticky="ew")

        # Test frame (bottom-left)
        frm_test = ttk.LabelFrame(self, text="Add / Edit Test")
        frm_test.place(x=10, y=250, width=460, height=330)

        ttk.Label(frm_test, text="Test Name:").grid(row=0, column=0, sticky="e", padx=6, pady=6)
        ttk.Label(frm_test, text="Date (YYYY-MM-DD):").grid(row=1, column=0, sticky="e", padx=6, pady=6)
        ttk.Label(frm_test, text="Result:").grid(row=2, column=0, sticky="ne", padx=6, pady=6)

        self.combo_test_name = ttk.Combobox(frm_test, values=TEST_CHOICES, state="readonly", width=33)
        self.combo_test_name.grid(row=0, column=1, padx=6, pady=6)
        self.combo_test_name.set(TEST_CHOICES[0])
        self.combo_test_name.bind("<<ComboboxSelected>>", self.on_test_choice_changed)

        self.e_test_date = ttk.Entry(frm_test, width=35)
        self.e_test_date.grid(row=1, column=1, padx=6, pady=6)
        if not self.e_test_date.get():
            self.e_test_date.insert(0, datetime.now().strftime("%Y-%m-%d"))

        self.e_result = tk.Text(frm_test, width=33, height=6)
        self.e_result.grid(row=2, column=1, padx=6, pady=6)

        # Buttons for test CRUD
        btn_add_test    = ttk.Button(frm_test, text="Add Test", command=self.add_test_btn)
        btn_update_test = ttk.Button(frm_test, text="Update Selected Test", command=self.update_test_btn)
        btn_delete_test = ttk.Button(frm_test, text="Delete Selected Test", command=self.delete_test_btn)
        btn_clear_test  = ttk.Button(frm_test, text="Clear Form", command=self.clear_test_form)

        btn_add_test.grid(row=3, column=0, padx=6, pady=6, sticky="ew")
        btn_update_test.grid(row=3, column=1, padx=6, pady=6, sticky="ew")
        btn_delete_test.grid(row=4, column=0, padx=6, pady=6, sticky="ew")
        btn_clear_test.grid(row=4, column=1, padx=6, pady=6, sticky="ew")

        # Tests list frame (right side) – wider now
        frm_list = ttk.LabelFrame(self, text="Patient Tests")
        frm_list.place(x=480, y=10, width=520, height=570)

        self.tests_tree = ttk.Treeview(
            frm_list,
            columns=("id","name","date","result","price"),
            show="headings",
            height=24
        )
        self.tests_tree.heading("id", text="ID")
        self.tests_tree.heading("name", text="Test")
        self.tests_tree.heading("date", text="Date")
        self.tests_tree.heading("result", text="Result")
        self.tests_tree.heading("price", text="Price (RM)")

        # Wider columns to fit price nicely
        self.tests_tree.column("id", width=50, anchor="center")
        self.tests_tree.column("name", width=170)
        self.tests_tree.column("date", width=100, anchor="center")
        self.tests_tree.column("result", width=150)
        self.tests_tree.column("price", width=90, anchor="e")

        self.tests_tree.pack(fill="both", expand=True, padx=6, pady=6)
        self.tests_tree.bind("<<TreeviewSelect>>", self.on_tree_select)

        btn_refresh = ttk.Button(frm_list, text="Refresh List", command=self.refresh_tests)
        btn_refresh.pack(padx=6, pady=(0,4), fill="x")

        # NEW: current price label (shows price for selected value in the test combobox)
        self.current_price_label = ttk.Label(frm_list, text="Current price: RM 0.00", anchor="e")
        self.current_price_label.pack(padx=6, pady=(0,8), fill="x")

        # Total payment label (large/bold)
        self.total_label = ttk.Label(frm_list, text="Total: RM 0.00", anchor="e")
        try:
            big = ("Segoe UI", 12, "bold")
        except:
            big = ("TkDefaultFont", 12, "bold")
        self.total_label.configure(font=big)
        self.total_label.pack(padx=6, pady=(0,8), fill="x")

        # Initialize current price display
        self.update_current_price_label()

    # ----- Prices Dialog -----
    def open_prices_dialog(self):
        prices = get_all_prices()

        top = tk.Toplevel(self)
        top.title("Set Test Prices")
        top.geometry("360x260")
        top.resizable(False, False)

        entries = {}
        row = 0
        for name in TEST_CHOICES:
            ttk.Label(top, text=name + " (RM)").grid(row=row, column=0, sticky="e", padx=8, pady=6)
            e = ttk.Entry(top, width=18)
            e.grid(row=row, column=1, padx=8, pady=6)
            e.insert(0, f"{prices.get(name, 0.0):.2f}")
            entries[name] = e
            row += 1

        def save_prices():
            new_prices = {}
            try:
                for name, e in entries.items():
                    val = float(e.get().strip())
                    if val < 0:
                        raise ValueError
                    new_prices[name] = val
            except ValueError:
                messagebox.showerror("Error", "Please enter valid non-negative numbers.")
                return
            set_all_prices(new_prices)
            messagebox.showinfo("Saved", "Prices updated.")
            top.destroy()
            # Update the current price label after saving
            self.update_current_price_label()
            # Refresh list total etc.
            self.refresh_tests()

        btns = ttk.Frame(top)
        btns.grid(row=row, column=0, columnspan=2, pady=10)
        ttk.Button(btns, text="Save", command=save_prices).pack(side="left", padx=8)
        ttk.Button(btns, text="Cancel", command=top.destroy).pack(side="left", padx=8)

    # ----- GUI actions -----
    def save_patient(self):
        b = self.e_barcode.get().strip()
        w = self.e_ward.get().strip()
        n = self.e_name.get().strip()
        ic = self.e_ic.get().strip()
        if not b:
            messagebox.showerror("Error", "Barcode is required.")
            return
        upsert_patient(b, w, n, ic)
        messagebox.showinfo("Saved", "Patient saved/updated.")

    def load_patient(self):
        b = self.e_barcode.get().strip()
        if not b:
            messagebox.showerror("Error", "Enter a barcode to load.")
            return
        rec = get_patient(b)
        if not rec:
            messagebox.showwarning("Not found", "No patient with that barcode.")
            return
        _, w, n, ic = rec
        self.e_ward.delete(0, tk.END); self.e_ward.insert(0, w or "")
        self.e_name.delete(0, tk.END); self.e_name.insert(0, n or "")
        self.e_ic.delete(0, tk.END); self.e_ic.insert(0, ic or "")
        self.refresh_tests()
        self.clear_test_form()

    def add_test_btn(self):
        b = self.e_barcode.get().strip()
        if not b:
            messagebox.showerror("Error", "Save/load a patient first.")
            return
        tname = self.combo_test_name.get().strip()
        tdate = self.e_test_date.get().strip()
        result = self.e_result.get("1.0", "end").strip()
        if not tname or not tdate:
            messagebox.showerror("Error", "Test name and date are required.")
            return
        # Snapshot price
        price = get_current_price(tname)
        add_test(b, tname, tdate, result, price)
        self.refresh_tests()
        self.clear_test_form()
        messagebox.showinfo("Added", f"Test added. Price snapshot: RM {price:.2f}")

    def on_tree_select(self, event):
        sel = self.tests_tree.selection()
        if not sel:
            self.selected_test_id = None
            return
        item = self.tests_tree.item(sel[0])
        test_id, tname, tdate, result, price = item["values"]
        self.selected_test_id = test_id
        # Fill form
        self.combo_test_name.set(tname)
        self.e_test_date.delete(0, tk.END); self.e_test_date.insert(0, tdate)
        self.e_result.delete("1.0", "end"); self.e_result.insert("1.0", result or "")
        # Also update the current price label to match selected test
        self.update_current_price_label()

    def update_test_btn(self):
        if not self.selected_test_id:
            messagebox.showerror("Error", "Select a test from the list to update.")
            return
        tname = self.combo_test_name.get().strip()
        tdate = self.e_test_date.get().strip()
        result = self.e_result.get("1.0", "end").strip()
        if not tname or not tdate:
            messagebox.showerror("Error", "Test name and date are required.")
            return
        # If test name changed, refresh price snapshot to current rate
        current_sel = self.tests_tree.selection()
        old_name = None
        if current_sel:
            vals = self.tests_tree.item(current_sel[0])["values"]
            if vals:
                old_name = vals[1]
        if old_name is not None and old_name != tname:
            new_price = get_current_price(tname)
            update_test(self.selected_test_id, tname, tdate, result, new_price)
            messagebox.showinfo("Updated", f"Test updated with new price snapshot: RM {new_price:.2f}")
        else:
            update_test(self.selected_test_id, tname, tdate, result, None)
            messagebox.showinfo("Updated", "Test updated.")
        self.refresh_tests()

    def delete_test_btn(self):
        if not self.selected_test_id:
            messagebox.showerror("Error", "Select a test from the list to delete.")
            return
        if not messagebox.askyesno("Confirm", "Delete the selected test?"):
            return
        delete_test(self.selected_test_id)
        self.refresh_tests()
        self.clear_test_form()
        messagebox.showinfo("Deleted", "Test deleted.")

    def clear_test_form(self):
        self.selected_test_id = None
        self.combo_test_name.set(TEST_CHOICES[0])
        self.e_test_date.delete(0, tk.END)
        self.e_test_date.insert(0, datetime.now().strftime("%Y-%m-%d"))
        self.e_result.delete("1.0", "end")
        # Update current price label to default combo choice
        self.update_current_price_label()

    def refresh_tests(self):
        for row in self.tests_tree.get_children():
            self.tests_tree.delete(row)
        b = self.e_barcode.get().strip()
        total = 0.0
        if b:
            rows = list_tests(b)
            for rid, name, date, result, price in rows:
                total += float(price or 0.0)
                self.tests_tree.insert("", "end", values=(rid, name, date, result, f"{price:.2f}"))
        # Update total label
        self.total_label.configure(text=f"Total: RM {total:.2f}")

    def print_sticker_btn(self):
        b = self.e_barcode.get().strip()
        w = self.e_ward.get().strip()
        n = self.e_name.get().strip()
        ic = self.e_ic.get().strip()
        if not b or not n:
            messagebox.showerror("Error", "Barcode and Name are required.")
            return
        print_sticker(b, w, n, ic)
        messagebox.showinfo("Print", "Sticker sent to printer.")

    def on_test_choice_changed(self, event=None):
        self.update_current_price_label()

    def update_current_price_label(self):
        tname = self.combo_test_name.get().strip()
        price = get_current_price(tname) if tname else 0.0
        self.current_price_label.configure(text=f"Current price: RM {price:.2f}")

# ========= main ========
if __name__ == "__main__":
    init_db()
    app = LISApp()
    app.mainloop()
