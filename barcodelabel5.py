import os
from datetime import datetime
from barcode import Code128
from barcode.writer import ImageWriter
from PIL import Image
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from tempfile import NamedTemporaryFile

# === Config ===
barcode_data = "941224085918"
line2_text   = "Rabiatul Adawiyah Binti Abdul Samad"
line3_text   = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

LABEL_W_CM = 5.5
LABEL_H_CM = 2.0
CROP_KEEP  = 0.30  # keep top 30% of barcode image

barcode_img     = "barcode.png"
barcode_cropped = "barcode_cropped.png"

# === Generate barcode and crop ===
barcode = Code128(barcode_data, writer=ImageWriter())
barcode.save(barcode_img.replace(".png", ""))  # writes "barcode.png"

img = Image.open(barcode_img)
w, h = img.size
cropped_h  = int(h * CROP_KEEP)
cropped_img = img.crop((0, 0, w, cropped_h))
cropped_img.save(barcode_cropped)

# === Create Word document (flush to top, no margins/spaces) ===
doc = Document()

# Kill hidden spacing on default Normal style
normal = doc.styles['Normal']
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after  = Pt(0)
normal.paragraph_format.line_spacing = 1.0

section = doc.sections[0]
section.page_width  = Cm(LABEL_W_CM)
section.page_height = Cm(LABEL_H_CM)
section.top_margin    = Cm(0)
section.bottom_margin = Cm(0)
section.left_margin   = Cm(0)
section.right_margin  = Cm(0)

# Single tightly-packed paragraph at the very top
p = doc.add_paragraph()
p.alignment = 1  # center horizontally
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after  = Pt(0)
p.paragraph_format.line_spacing = 1.0

# === Barcode image (explicit small height; adjust if you want more/less) ===
run = p.add_run()
run.add_picture(barcode_cropped, width=Cm(LABEL_W_CM), height=Cm(0.7))  # 0.7 cm tall

# === First line (barcode number) ===
run.add_break()
t1 = p.add_run(barcode_data)
t1.font.name = "Calibri"
t1.font.size = Pt(8)
t1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

# === Second line (name) ===
t1.add_break()
t2 = p.add_run(line2_text)
t2.font.name = "Calibri"
t2.font.size = Pt(8)
t2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

# === Third line (date/time) ===
t2.add_break()
t3 = p.add_run(line3_text)
t3.font.name = "Calibri"
t3.font.size = Pt(8)
t3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

# === Save and print ===
with NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
    doc_path = tmp_file.name
    doc.save(doc_path)

os.startfile(doc_path, "print")
